import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
from transformers import pipeline
from docx import Document
import fitz  # for reading PDFs
import os

loaded_file_extension = ""
# Initialize the summarization pipeline
summarizer = pipeline("summarization", model="sshleifer/distilbart-cnn-12-6")

def load_text(file_path):
    if file_path.endswith('.txt'):
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
    elif file_path.endswith('.docx'):
        doc = Document(file_path)
        text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
    elif file_path.endswith('.pdf'):
        try:
            doc = fitz.open(file_path)
            text = ""
            for page in doc:
                text += page.get_text()
        except Exception as e:
            messagebox.showerror("Loading Error", f"Failed to load PDF: {e}")
            return ""
    else:
        messagebox.showerror("Unsupported Format", "This file format is not supported.")
        return ""
    return text

def load_text_file():
    file_path = filedialog.askopenfilename(filetypes=[("Text Files", "*.txt"), 
                                                       ("Word Documents", "*.docx"), 
                                                       ("PDF Files", "*.pdf")])
    if file_path:
        text = load_text(file_path)
        text_area.delete('1.0', tk.END)
        text_area.insert(tk.END, text)
        # Store the loaded file's extension for later use
        global loaded_file_extension
        loaded_file_extension = file_path.split('.')[-1]
def load_folder():
    folder_path = filedialog.askdirectory()
    if folder_path:
        document_count = 0
        for filename in os.listdir(folder_path):
            file_path = os.path.join(folder_path, filename)
            if filename.endswith(('.txt', '.docx', '.pdf')):
                text = load_text(file_path)
                if text:
                    summary = summarizer(text, max_length=75, min_length=40, do_sample=False)
                    summary_text = summary[0]['summary_text']
                    # Append each summary to the summary_area
                    summary_area.insert(tk.END, f"Summary of {filename}:\n{summary_text}\n\n")
                    document_count += 1
        
        if document_count > 0:
            messagebox.showinfo("Success", f"Folder loaded successfully. {document_count} documents processed.")
        else:
            messagebox.showinfo("No Documents Found", "No supported documents found in the selected folder.")

        
def save_summary():
    summary_text = summary_area.get('1.0', tk.END).strip()
    if not summary_text:
        messagebox.showerror("Error", "There is no summary to save.")
        return
    file_path = filedialog.asksaveasfilename(defaultextension=f".{loaded_file_extension}",
                                             filetypes=[("Text Files", "*.txt"),
                                                        ("Word Documents", "*.docx"),
                                                        ("PDF Files", "*.pdf")])
    if file_path:
        if file_path.endswith('.txt'):
            with open(file_path, 'w', encoding='utf-8') as file:
                file.write(summary_text)
        elif file_path.endswith('.docx'):
            doc = Document()
            doc.add_paragraph(summary_text)
            doc.save(file_path)
        messagebox.showinfo("Success", "The summary has been saved successfully.")

def generate_summary():
    text = text_area.get('1.0', tk.END)
    if text.strip():
        summary = summarizer(text, max_length=75, min_length=40, do_sample=False)
        summary_area.delete('1.0', tk.END)
        summary_area.insert(tk.END, summary[0]['summary_text'])
    else:
        messagebox.showerror("Error", "Please load a text file or folder to generate summaries.")

# Set up the GUI
root = tk.Tk()
root.title("Text Summarizer")

# Large title with deep pink background
title = tk.Label(root, text="Text Summarizer", font=("Arial", 30), bg="DeepPink3", fg="white")
title.pack(pady=20)

load_file_button = tk.Button(root, text="Load Text File", command=load_text_file, bg="blue", fg="white")
load_file_button.pack(pady=10)

load_folder_button = tk.Button(root, text="Load Folder", command=load_folder, bg="purple", fg="white")
load_folder_button.pack(pady=10)

text_area = scrolledtext.ScrolledText(root, height=10, width=100)
text_area.pack(pady=5)

summarize_button = tk.Button(root, text="Generate Summary", command=generate_summary, bg="yellow")
summarize_button.pack(pady=5)

summary_area = scrolledtext.ScrolledText(root, height=10, width=100)
summary_area.pack(pady=5)

save_button = tk.Button(root, text="Save Summary", command=save_summary, bg="green")
save_button.pack(pady=5)

root.mainloop()



"""
Text Summarizer Application

Description:
The Text Summarizer is a Python-based application that uses deep learning to automatically summarize text 
from .txt, .docx, and .pdf files. It enables users to either upload individual files or process all supported 
documents within a selected folder, displaying the summaries within an intuitive GUI and allowing the summaries 
to be saved in the original file format.

Features:
- Summarizes content from .txt, .docx, and .pdf files.
- Batch processing capability for handling multiple documents within a folder.
- An intuitive graphical user interface for easy interaction with the application.
- Save functionality enables users to save the summarized content back to the file system.

Installation and Setup:

1. Ensure Python 3.x is installed on your system.

2. Setup the Virtual Environment:
   - Navigate to the project directory.
   - Create the virtual environment by running: python -m venv myenv
   - Activate the virtual environment:
     - On Windows, run: myenv\\Scripts\\activate
     

3. Install Required Libraries:
   - Ensure the virtual environment is activated.
   - Install the necessary libraries by running: pip install transformers python-docx PyMuPDF

"""
